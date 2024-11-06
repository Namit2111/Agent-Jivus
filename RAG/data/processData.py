from docx import Document
import re
import csv
import json
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
def load_pdf(filename):
    loader = PyPDFLoader(filename)
    pages = loader.load_and_split()
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)

    documents = text_splitter.split_documents(pages)

    return documents

def load_docx(file_path): #notused
    """Loads content from a .docx file and returns it as a string."""
    doc = Document(file_path)
    content = ""
    for para in doc.paragraphs:
        if para.text.strip():  
            content += para.text + "\n"
    return content.strip()

def chunk_text_with_overlap(text, chunk_size=128, overlap_size=20): #notused
    """Chunk text with overlap to prevent loss of context between chunks."""
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap_size):
        chunk = ' '.join(words[i:i + chunk_size])
        chunks.append(chunk)
    return chunks

def format_product_details(documents):
    product_details=[]
    for doc in documents:
        product_details.append(
            {
                "page_content":doc.page_content,
                "metadata":{
                    "category":"Product Details",
                    "page_number":doc.metadata['page'] + 1
                }
            }
        )
    return product_details

def split_dialogues_to_csv(doc_path, csv_path):
    
    doc = Document(doc_path)
    
    conversation_start_pattern = re.compile(r"^Conversation\s(\d+):", re.IGNORECASE)
    dialogue_pattern = re.compile(r"^(Prospect|SDR):\s*(.*)", re.IGNORECASE)
    
    conversation_id = None
    rows = []

    for para in doc.paragraphs:
        text = para.text.strip()
        
        conversation_start_match = conversation_start_pattern.match(text)
        if conversation_start_match:
            conversation_id = conversation_start_match.group(1) 
            continue  
        
        dialogue_match = dialogue_pattern.match(text)
        if dialogue_match and conversation_id:
            person = dialogue_match.group(1) 
            dialogue = dialogue_match.group(2) 
            rows.append([conversation_id, person, dialogue])

    with open(csv_path, mode='w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(["convoid", "person", "dialogue"]) 
        writer.writerows(rows)



def convert_csv_to_conversation_data(csv_path):
    conversation_data = []
    previous_row = None  # To keep track of the previous dialogue for linking

    with open(csv_path, mode="r", encoding="utf-8") as file:
        reader = list(csv.DictReader(file))  # Convert reader to list for indexing

        for i, row in enumerate(reader):
            # Extract data from each row
            convoid = row['convoid']
            person = row['person']
            dialogue = row['dialogue']
            
            # Determine the next dialogue, if it exists and is in the same conversation
            next_dialogue = (
                f"{reader[i + 1]['person']}: {reader[i + 1]['dialogue']}"
                if i + 1 < len(reader) and reader[i + 1]['convoid'] == convoid
                else ""
            )

            # Create an entry with the current dialogue and link the next dialogue as an answer
            conversation_data.append({
                "page_content": f"{person}: {dialogue}",
                "metadata": {
                    "category": "Conversation",
                    "conversation_id": convoid,
                    "answer": next_dialogue if next_dialogue else ""
                }
            })
    return conversation_data



def save_conversations_to_json(doc_path, json_path):
    doc = Document(doc_path)
    
    # Regular expressions for detecting conversation start and dialogue lines
    conversation_start_pattern = re.compile(r"^Conversation\s(\d+):", re.IGNORECASE)
    dialogue_pattern = re.compile(r"^(Prospect|SDR):\s*(.*)", re.IGNORECASE)
    
    conversations = []
    conversation_id = None
    current_conversation = []

    # Iterate through each paragraph in the document
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Detect start of a new conversation
        conversation_start_match = conversation_start_pattern.match(text)
        if conversation_start_match:
            # Save the current conversation if it exists
            if current_conversation:
                conversations.append({
                    "page_content": " ".join(current_conversation),
                    "metadata": {
                        "category": "Conversation",
                        "conversation_id": conversation_id
                    }
                })
                current_conversation = []

            # Update conversation ID for the new conversation
            conversation_id = conversation_start_match.group(1)
            continue

        # Detect dialogue lines within a conversation
        dialogue_match = dialogue_pattern.match(text)
        if dialogue_match and conversation_id:
            person = dialogue_match.group(1)
            dialogue = dialogue_match.group(2)
            current_conversation.append(f"{person}: {dialogue}")

    # Add the last conversation if there's any content left
    if current_conversation:
        conversations.append({
            "page_content": " ".join(current_conversation),
            "metadata": {
                "category": "Conversation",
                "conversation_id": conversation_id
            }
        })

    # Write the conversations to the JSON file
    with open(json_path, mode='w', encoding='utf-8') as file:
        json.dump(conversations, file, indent=1)


if __name__ =="__main__":

    documents = load_pdf(filename="./productDetails.pdf")
    pdata = format_product_details(documents)
    with open("productDetails.json",'w') as f:
        json.dump(pdata,f,indent=1)
    doc_path = "testConvo.docx" 
    csv_path = "conversations.csv"  
    split_dialogues_to_csv(doc_path, csv_path)
    data = convert_csv_to_conversation_data(csv_path="./conversations.csv")
    with open("convoDetails.json",'w') as d:
        json.dump(data,d,indent=1)

    save_conversations_to_json(doc_path=doc_path, json_path="wholeConvo.json")
